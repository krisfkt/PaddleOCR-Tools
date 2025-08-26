# paddleocr_fixed.py - AI-Generated PaddleOCR Testing Demo Tools
# 完整版配置化 OCR 處理器
import os
import time
import json
import argparse
import configparser
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Tuple, Optional

# 核心套件
import cv2
import numpy as np
from PIL import Image, ImageDraw, ImageFont
from paddleocr import PaddleOCR

# 文件處理
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# PDF 生成
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


class PaddleOCRResultParser:
    """PaddleOCR 結果解析器 - 適配新版 API"""
    
    @staticmethod
    def extract_text_from_result(result) -> Tuple[str, List[Dict]]:
        """從新版 PaddleOCR 結果中提取文字"""
        text_lines = []
        detailed_results = []
        
        try:
            if isinstance(result, list) and len(result) > 0:
                # 新版返回 OCRResult 物件
                ocr_result = result[0]
                
                # 檢查是否有 rec_texts 屬性（識別的文字）
                if hasattr(ocr_result, 'rec_texts') and ocr_result.rec_texts:
                    texts = ocr_result.rec_texts
                    scores = getattr(ocr_result, 'rec_scores', [1.0] * len(texts))
                    polys = getattr(ocr_result, 'rec_polys', [None] * len(texts))
                    
                    for i, text in enumerate(texts):
                        confidence = scores[i] if i < len(scores) else 1.0
                        poly = polys[i] if i < len(polys) else None
                        
                        text_lines.append(text)
                        detailed_results.append({
                            'text': text,
                            'confidence': confidence,
                            'bbox': poly.tolist() if poly is not None else None
                        })
                
                # 如果是字典格式（向後兼容）
                elif isinstance(ocr_result, dict):
                    if 'rec_texts' in ocr_result and ocr_result['rec_texts']:
                        texts = ocr_result['rec_texts']
                        scores = ocr_result.get('rec_scores', [1.0] * len(texts))
                        polys = ocr_result.get('rec_polys', [None] * len(texts))
                        
                        for i, text in enumerate(texts):
                            confidence = scores[i] if i < len(scores) else 1.0
                            poly = polys[i] if i < len(polys) else None
                            
                            text_lines.append(text)
                            detailed_results.append({
                                'text': text,
                                'confidence': confidence,
                                'bbox': poly.tolist() if poly is not None else None
                            })
        
        except Exception as e:
            print(f"結果解析錯誤: {e}")
        
        return '\n'.join(text_lines), detailed_results


class OptimizedOCRProcessor:
    """AI-Generated PaddleOCR Testing Demo Tools - 核心處理器"""
    
    def __init__(self, config_file: str = "paddleocr_config.ini"):
        self.config_file = config_file
        self.ocr_engine = None
        self.parser = PaddleOCRResultParser()
        self.config = self._load_config()
        self._init_ocr_engine()
        print("✓ AI-Generated PaddleOCR Testing Demo Tools 已初始化")
    
    def _load_config(self):
        """載入配置檔案"""
        config = configparser.ConfigParser()
        
        # 預設配置
        default_config = {
            'OCR': {
                'lang': 'chinese_cht',
                'use_angle_cls': 'True',
                'use_gpu': 'False',
                'show_log': 'False'
            },
            'PROCESSING': {
                'confidence_threshold': '0.9',
                'default_output_format': 'pdf'
            },
            'OUTPUT': {
                'output_folder': './output',
                'include_stats': 'False',
                'save_raw_results': 'False',
                'simple_output': 'True'
            }
        }
        
        # 如果配置檔案存在，讀取它
        if os.path.exists(self.config_file):
            try:
                config.read(self.config_file, encoding='utf-8')
                print(f"✓ 載入配置檔案: {self.config_file}")
            except Exception as e:
                print(f"✗ 配置檔案讀取失敗: {e}")
                print("使用預設配置")
        else:
            # 創建預設配置檔案
            for section, options in default_config.items():
                config.add_section(section)
                for key, value in options.items():
                    config.set(section, key, value)
            
            try:
                with open(self.config_file, 'w', encoding='utf-8') as f:
                    config.write(f)
                print(f"✓ 創建預設配置檔案: {self.config_file}")
            except Exception as e:
                print(f"✗ 配置檔案創建失敗: {e}")
        
        return config
    
    def get_config_value(self, section: str, key: str, fallback=None, value_type=str):
        """安全獲取配置值"""
        try:
            if value_type == bool:
                return self.config.getboolean(section, key, fallback=fallback)
            elif value_type == float:
                return self.config.getfloat(section, key, fallback=fallback)
            elif value_type == int:
                return self.config.getint(section, key, fallback=fallback)
            else:
                return self.config.get(section, key, fallback=fallback)
        except:
            return fallback
    
    def _init_ocr_engine(self):
        """根據配置初始化 OCR 引擎"""
        try:
            print("正在初始化 PaddleOCR 引擎...")
            
            # 從配置檔案獲取參數
            ocr_config = {}
            lang = self.get_config_value('OCR', 'lang', 'chinese_cht')
            
            # 語言映射和驗證
            supported_langs = {
                'ch': '簡體中文',
                'chinese_cht': '繁體中文', 
                'en': '英文',
                'chinese_traditional': '繁體中文',
                'cht': '繁體中文'
            }
            
            # 標準化語言代碼
            if lang in ['chinese_cht', 'chinese_traditional', 'cht']:
                ocr_config['lang'] = 'chinese_cht'
                print(f"✓ 使用繁體中文模型")
            elif lang == 'ch':
                ocr_config['lang'] = 'ch'
                print(f"✓ 使用簡體中文模型")
            elif lang == 'en':
                ocr_config['lang'] = 'en'
                print(f"✓ 使用英文模型")
            else:
                print(f"⚠️ 不支援的語言: {lang}，使用預設繁體中文")
                ocr_config['lang'] = 'chinese_cht'
            
            ocr_config['use_angle_cls'] = self.get_config_value('OCR', 'use_angle_cls', True, bool)
            ocr_config['use_gpu'] = self.get_config_value('OCR', 'use_gpu', False, bool)
            
            # show_log 在新版本可能不支援，所以先嘗試
            try:
                show_log = self.get_config_value('OCR', 'show_log', False, bool)
                ocr_config['show_log'] = show_log
            except:
                pass
            
            print(f"使用配置: {ocr_config}")
            
            # 備用配置（如果配置檔案失敗）
            configs_to_try = [
                ocr_config,
                {'lang': 'chinese_cht'},
                {'lang': 'ch'},
                {'lang': 'en'},
            ]
            
            for i, config in enumerate(configs_to_try):
                try:
                    print(f"嘗試配置 {i+1}: {config}")
                    self.ocr_engine = PaddleOCR(**config)
                    self.current_config = config
                    print(f"✓ 成功初始化，配置: {config}")
                    
                    # 顯示當前使用的語言
                    used_lang = config.get('lang', 'unknown')
                    if used_lang in supported_langs:
                        print(f"✓ 當前語言模型: {supported_langs[used_lang]}")
                    
                    break
                except Exception as e:
                    print(f"✗ 配置失敗: {e}")
                    if 'chinese_cht' in str(config) and 'not found' in str(e).lower():
                        print("  💡 提示：可能需要下載繁體中文模型")
            
            if self.ocr_engine is None:
                raise Exception("所有配置都失敗")
            
        except Exception as e:
            print(f"OCR 初始化失敗: {e}")
            self.ocr_engine = None
    
    def process_image(self, image_path: str, confidence_threshold: float = None) -> Dict:
        """處理圖片檔案"""
        if not os.path.exists(image_path):
            raise FileNotFoundError(f"找不到圖片檔案: {image_path}")
        
        if self.ocr_engine is None:
            raise RuntimeError("PaddleOCR 引擎未初始化")
        
        # 如果沒有指定信心度閾值，使用配置檔案中的值
        if confidence_threshold is None:
            confidence_threshold = self.get_config_value('PROCESSING', 'confidence_threshold', 0.9, float)
        
        # 檢查是否為簡單輸出模式
        simple_output = self.get_config_value('OUTPUT', 'simple_output', True, bool)
        
        if not simple_output:
            print(f"開始處理圖片: {image_path}")
            print(f"使用信心度閾值: {confidence_threshold}")
        
        start_time = time.time()
        
        try:
            # 載入圖片
            image = cv2.imread(image_path)
            if image is None:
                pil_image = Image.open(image_path)
                image = np.array(pil_image)
                if len(image.shape) == 3:
                    image = cv2.cvtColor(image, cv2.COLOR_RGB2BGR)
            
            # 執行 OCR
            self._print_if_verbose("正在執行 OCR 識別...")
            
            try:
                results = self.ocr_engine.predict(image)
            except AttributeError:
                results = self.ocr_engine.ocr(image)
            
            end_time = time.time()
            processing_time = end_time - start_time
            
            # 解析結果
            text_content, detailed_results = self.parser.extract_text_from_result(results)
            
            # 過濾低信心度結果
            filtered_results = []
            filtered_text_lines = []
            
            for result in detailed_results:
                if result['confidence'] >= confidence_threshold:
                    filtered_results.append(result)
                    filtered_text_lines.append(result['text'])
            
            filtered_text = '\n'.join(filtered_text_lines)
            
            # 統計資訊
            stats = {
                'processing_time': f"{processing_time:.2f} 秒",
                'total_detected': len(detailed_results),
                'accepted_lines': len(filtered_results),
                'total_chars': len(filtered_text),
                'total_words': len(filtered_text.split()),
                'confidence_threshold': confidence_threshold,
                'average_confidence': sum(r['confidence'] for r in filtered_results) / len(filtered_results) if filtered_results else 0
            }
            
            # 根據模式顯示結果
            if simple_output:
                # 只顯示 OCR 結果
                self._print_ocr_result_only({'text_content': filtered_text})
            else:
                # 顯示詳細信息
                print(f"✓ OCR 完成")
                print(f"  檢測到: {stats['total_detected']} 行")
                print(f"  接受: {stats['accepted_lines']} 行")
                print(f"  字符數: {stats['total_chars']}")
                print(f"  平均信心度: {stats['average_confidence']:.3f}")
            
            return {
                'text_content': filtered_text,
                'all_text': text_content,
                'detailed_results': filtered_results,
                'all_results': detailed_results,
                'stats': stats,
                'raw_ocr_result': results if self.get_config_value('OUTPUT', 'save_raw_results', False, bool) else None
            }
            
        except Exception as e:
            raise Exception(f"圖片處理失敗: {e}")
        
    def _print_if_verbose(self, message):
        """根據配置決定是否顯示訊息"""
        simple_output = self.get_config_value('OUTPUT', 'simple_output', True, bool)
        if not simple_output:
            print(message)

    def _print_ocr_result_only(self, result):
        """只顯示 OCR 結果"""
        if result['text_content'].strip():
            print(result['text_content'])
        else:
            print("(無識別內容)")

    def save_result_to_file(self, result: Dict, input_file: str, format_type: str = None) -> str:
        """儲存結果到檔案"""
        # 如果沒有指定格式，使用配置檔案中的預設格式
        if format_type is None:
            format_type = self.get_config_value('PROCESSING', 'default_output_format', 'pdf')
        
        # 獲取輸出資料夾
        output_folder = self.get_config_value('OUTPUT', 'output_folder', './output')
        os.makedirs(output_folder, exist_ok=True)
        
        # 是否包含統計資訊
        include_stats = self.get_config_value('OUTPUT', 'include_stats', False, bool)
        
        # 是否為簡單輸出模式
        simple_output = self.get_config_value('OUTPUT', 'simple_output', True, bool)
        
        base_name = Path(input_file).stem
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if format_type == 'txt':
            # 簡單模式下的檔名
            if simple_output:
                output_path = os.path.join(output_folder, f"{base_name}_ocr.txt")
            else:
                output_path = os.path.join(output_folder, f"{base_name}_{timestamp}_fixed.txt")
            
            with open(output_path, 'w', encoding='utf-8') as f:
                if simple_output:
                    # 簡單模式：只寫入 OCR 結果
                    f.write(result['text_content'])
                else:
                    # 詳細模式：包含完整資訊
                    f.write("=== AI-Generated PaddleOCR Testing Demo Tools 處理結果 ===\n")
                    f.write(f"原始檔案: {Path(input_file).name}\n")
                    f.write(f"處理時間: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                    f.write(f"配置檔案: {self.config_file}\n\n")
                    
                    if include_stats:
                        stats = result['stats']
                        f.write("=== 處理統計 ===\n")
                        f.write(f"處理時間: {stats['processing_time']}\n")
                        f.write(f"檢測行數: {stats['total_detected']}\n")
                        f.write(f"接受行數: {stats['accepted_lines']}\n")
                        f.write(f"字符數: {stats['total_chars']}\n")
                        f.write(f"詞數: {stats['total_words']}\n")
                        f.write(f"信心度閾值: {stats['confidence_threshold']}\n")
                        f.write(f"平均信心度: {stats['average_confidence']:.3f}\n\n")
                    
                    f.write("=== 識別內容（過濾後）===\n")
                    f.write(result['text_content'])
                    f.write("\n\n")
                    
                    if include_stats:
                        f.write("=== 詳細結果 ===\n")
                        for i, item in enumerate(result['detailed_results']):
                            f.write(f"行 {i+1}: '{item['text']}' (信心度: {item['confidence']:.3f})\n")
                            if item['bbox']:
                                f.write(f"  座標: {item['bbox']}\n")
                            f.write("\n")
                    
                    if result['all_results'] != result['detailed_results'] and include_stats:
                        f.write("=== 所有檢測結果（包含低信心度）===\n")
                        for i, item in enumerate(result['all_results']):
                            f.write(f"行 {i+1}: '{item['text']}' (信心度: {item['confidence']:.3f})\n")
                    
                    # 保存原始 OCR 結果（如果啟用）
                    if result['raw_ocr_result'] is not None:
                        f.write("\n=== 原始 OCR 結果 ===\n")
                        f.write(json.dumps(str(result['raw_ocr_result']), ensure_ascii=False, indent=2))
            
            if not simple_output:
                print(f"✓ 結果已儲存: {output_path}")
            return output_path
        
        elif format_type == 'docx':
            # 簡單模式下的檔名
            if simple_output:
                output_path = os.path.join(output_folder, f"{base_name}_ocr.docx")
            else:
                output_path = os.path.join(output_folder, f"{base_name}_{timestamp}_fixed.docx")
            
            doc = docx.Document()
            
            # 設定字體
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Microsoft YaHei'
            font.size = Pt(12)
            
            if simple_output:
                # 簡單模式：只包含 OCR 內容
                paragraphs = result['text_content'].split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        doc.add_paragraph(para_text)
                    else:
                        doc.add_paragraph()
            else:
                # 詳細模式：包含完整資訊
                title = doc.add_heading('AI-Generated PaddleOCR Testing Demo Tools 識別結果', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 檔案資訊
                info_para = doc.add_paragraph()
                info_para.add_run(f"原始檔案: {Path(input_file).name}\n").bold = True
                info_para.add_run(f"處理時間: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                info_para.add_run(f"配置檔案: {self.config_file}\n")
                
                # 統計資訊（如果啟用）
                if include_stats:
                    stats = result['stats']
                    stats_para = doc.add_paragraph()
                    stats_para.add_run("處理統計:\n").bold = True
                    stats_para.add_run(f"• 處理時間: {stats['processing_time']}\n")
                    stats_para.add_run(f"• 檢測行數: {stats['total_detected']}\n")
                    stats_para.add_run(f"• 接受行數: {stats['accepted_lines']}\n")
                    stats_para.add_run(f"• 字符數: {stats['total_chars']}\n")
                    stats_para.add_run(f"• 信心度閾值: {stats['confidence_threshold']}\n")
                    stats_para.add_run(f"• 平均信心度: {stats['average_confidence']:.3f}\n")
                
                # 分隔線
                doc.add_paragraph("=" * 50)
                
                # 識別內容
                content_heading = doc.add_heading('識別內容', level=1)
                
                paragraphs = result['text_content'].split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        doc.add_paragraph(para_text)
                    else:
                        doc.add_paragraph()
            
            doc.save(output_path)
            if not simple_output:
                print(f"✓ DOCX 檔案已儲存: {output_path}")
            return output_path
        
        elif format_type == 'pdf':
            # 簡單模式下的檔名
            if simple_output:
                output_path = os.path.join(output_folder, f"{base_name}_ocr.pdf")
            else:
                output_path = os.path.join(output_folder, f"{base_name}_{timestamp}_fixed.pdf")
            
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            
            # 嘗試註冊中文字體
            try:
                font_paths = [
                    "C:/Windows/Fonts/msjh.ttc",    # 微軟正黑體（繁體）
                    "C:/Windows/Fonts/msyh.ttc",    # 微軟雅黑（簡體）
                    "C:/Windows/Fonts/simsun.ttc",  # 宋體
                ]
                for font_path in font_paths:
                    if os.path.exists(font_path):
                        pdfmetrics.registerFont(TTFont('Chinese', font_path))
                        styles['Normal'].fontName = 'Chinese'
                        styles['Title'].fontName = 'Chinese'
                        styles['Heading1'].fontName = 'Chinese'
                        break
            except:
                pass
            
            story = []
            
            if simple_output:
                # 簡單模式：只包含 OCR 內容
                paragraphs = result['text_content'].split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        safe_text = para_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        para = Paragraph(safe_text, styles['Normal'])
                        story.append(para)
                        story.append(Spacer(1, 6))
            else:
                # 詳細模式：包含完整資訊
                title = Paragraph("AI-Generated PaddleOCR Testing Demo Tools 識別結果", styles['Title'])
                story.append(title)
                story.append(Spacer(1, 20))
                
                # 檔案資訊
                info_text = f"""
                原始檔案: {Path(input_file).name}<br/>
                處理時間: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}<br/>
                配置檔案: {self.config_file}
                """
                
                if include_stats:
                    stats = result['stats']
                    info_text += f"""<br/>
                    檢測行數: {stats['total_detected']}<br/>
                    接受行數: {stats['accepted_lines']}<br/>
                    字符數: {stats['total_chars']}<br/>
                    信心度閾值: {stats['confidence_threshold']}<br/>
                    平均信心度: {stats['average_confidence']:.3f}
                    """
                
                info_para = Paragraph(info_text, styles['Normal'])
                story.append(info_para)
                story.append(Spacer(1, 20))
                
                # 內容標題
                content_title = Paragraph("識別內容", styles['Heading1'])
                story.append(content_title)
                story.append(Spacer(1, 12))
                
                # 內容
                paragraphs = result['text_content'].split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        safe_text = para_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        para = Paragraph(safe_text, styles['Normal'])
                        story.append(para)
                        story.append(Spacer(1, 6))
            
            doc.build(story)
            if not simple_output:
                print(f"✓ PDF 檔案已儲存: {output_path}")
            return output_path

    def create_better_test_image(self, text: str = "測試繁體中文 English 123", output_path: str = "ai_demo_test.png"):
        """建立 AI-Generated 測試圖片"""
        img = Image.new('RGB', (1200, 400), 'white')
        draw = ImageDraw.Draw(img)
        
        try:
            # 嘗試使用支援繁體中文的字體
            font_paths = [
                "C:/Windows/Fonts/msjh.ttc",    # 微軟正黑體（繁體）
                "C:/Windows/Fonts/msyh.ttc",    # 微軟雅黑
                "C:/Windows/Fonts/simsun.ttc",  # 宋體
                "C:/Windows/Fonts/arial.ttf",   # Arial
            ]
            
            font = None
            for font_path in font_paths:
                try:
                    font = ImageFont.truetype(font_path, 48)
                    print(f"✓ 使用字體: {font_path}")
                    break
                except:
                    continue
            
            if font is None:
                font = ImageFont.load_default()
                print("⚠️ 使用預設字體")
        
        except:
            font = ImageFont.load_default()
        
        # 計算文字位置（置中）
        bbox = draw.textbbox((0, 0), text, font=font)
        text_width = bbox[2] - bbox[0]
        text_height = bbox[3] - bbox[1]
        
        x = (img.width - text_width) // 2
        y = (img.height - text_height) // 2
        
        # 添加背景增強對比度
        padding = 30
        draw.rectangle([x-padding, y-padding//2, x+text_width+padding, y+text_height+padding//2], 
                      fill='white', outline='gray', width=3)
        
        # 繪製文字
        draw.text((x, y), text, font=font, fill='black')
        
        # 添加標籤
        label_font = ImageFont.load_default()
        label_text = "AI-Generated PaddleOCR Testing Demo"
        draw.text((10, 10), label_text, font=label_font, fill='gray')
        
        img.save(output_path, 'PNG', quality=95)
        print(f"✓ AI-Generated 測試圖片已建立: {output_path}")
        return output_path
    
    def batch_process_images(self, image_folder: str, confidence_threshold: float = None, output_format: str = None) -> List[str]:
        """批量處理圖片"""
        if not os.path.exists(image_folder):
            raise FileNotFoundError(f"找不到資料夾: {image_folder}")
        
        # 使用配置檔案的預設值
        if confidence_threshold is None:
            confidence_threshold = self.get_config_value('PROCESSING', 'confidence_threshold', 0.9, float)
        
        if output_format is None:
            output_format = self.get_config_value('PROCESSING', 'default_output_format', 'pdf')
        
        simple_output = self.get_config_value('OUTPUT', 'simple_output', True, bool)
        
        # 支援的圖片格式
        supported_formats = ('.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif')
        image_files = []
        
        for file in os.listdir(image_folder):
            if file.lower().endswith(supported_formats):
                image_files.append(os.path.join(image_folder, file))
        
        if not image_files:
            self._print_if_verbose("資料夾中沒有找到支援的圖片檔案")
            return []
        
        if not simple_output:
            print(f"AI-Generated PaddleOCR Testing Demo - 批量處理")
            print(f"找到 {len(image_files)} 個圖片檔案")
            print(f"使用信心度閾值: {confidence_threshold}")
            print(f"輸出格式: {output_format}")
        
        output_files = []
        
        for i, image_file in enumerate(image_files, 1):
            if simple_output:
                # 簡單模式：只顯示檔案名和結果
                print(f"=== {os.path.basename(image_file)} ===")
            else:
                # 詳細模式
                print(f"\n處理第 {i}/{len(image_files)} 個檔案: {os.path.basename(image_file)}")
            
            try:
                result = self.process_image(image_file, confidence_threshold)
                output_file = self.save_result_to_file(result, image_file, output_format)
                output_files.append(output_file)
                
                if not simple_output:
                    print(f"✓ 完成，識別 {len(result['text_content'])} 個字符")
                    
            except Exception as e:
                if simple_output:
                    print("(處理失敗)")
                else:
                    print(f"✗ 處理失敗: {e}")
        
        if not simple_output:
            print(f"\n批量處理完成，成功處理 {len(output_files)}/{len(image_files)} 個檔案")
        
        return output_files
    
    def print_current_config(self):
        """顯示當前配置"""
        print("\n=== AI-Generated PaddleOCR Testing Demo Tools - 當前配置 ===")
        print(f"配置檔案: {self.config_file}")
        
        for section in self.config.sections():
            print(f"\n[{section}]")
            for key, value in self.config.items(section):
                print(f"  {key} = {value}")


def main():
    """主函數"""
    print("🤖 AI-Generated PaddleOCR Testing Demo Tools")
    print("=" * 50)
    
    parser = argparse.ArgumentParser(description='AI-Generated PaddleOCR Testing Demo Tools - 配置版 OCR 處理器')
    parser.add_argument('--file', help='圖片檔案路徑')
    parser.add_argument('--folder', help='圖片資料夾路徑（批量處理）')
    parser.add_argument('--format', choices=['txt', 'docx', 'pdf'], help='輸出格式（覆蓋配置檔案）')
    parser.add_argument('--confidence', type=float, help='信心度閾值（覆蓋配置檔案）')
    parser.add_argument('--lang', choices=['ch', 'chinese_cht', 'en'], help='語言模型（覆蓋配置檔案）')
    parser.add_argument('--test', action='store_true', help='建立 AI-Generated 測試圖片並處理')
    parser.add_argument('--config', help='指定配置檔案路徑')
    parser.add_argument('--show-config', action='store_true', help='顯示當前配置')
    parser.add_argument('--simple', action='store_true', help='簡單輸出模式（只顯示 OCR 結果）')
    
    args = parser.parse_args()
    
    # 使用指定的配置檔案
    config_file = args.config if args.config else "paddleocr_config.ini"
    processor = OptimizedOCRProcessor(config_file)
    
    # 如果命令列指定了語言，臨時覆蓋配置
    if args.lang:
        processor.config.set('OCR', 'lang', args.lang)
        print(f"✓ 命令列覆蓋語言設定: {args.lang}")
        # 重新初始化 OCR 引擎
        processor._init_ocr_engine()
        
    # 如果命令列指定了簡單模式，臨時覆蓋配置
    if args.simple:
        processor.config.set('OUTPUT', 'simple_output', 'True')
    
    if processor.ocr_engine is None:
        print("OCR 引擎初始化失敗")
        return
    
    if args.show_config:
        processor.print_current_config()
        return
    
    if args.test:
        print("🧪 建立並測試 AI-Generated 示例圖片...")
        
        test_images = [
            ("Hello World", "ai_demo_test_en.png"),
            ("測試繁體中文", "ai_demo_test_cht.png"),
            ("發票INVOICE\n發票號碼InvoiceNo:AB-12345678\n客戶Customer：張三", "ai_demo_invoice.png"),
            ("混合 Mixed 123", "ai_demo_mixed.png")
        ]
        
        for text, filename in test_images:
            print(f"\n🔍 處理測試: {text[:20]}...")
            img_path = processor.create_better_test_image(text, filename)
            
            try:
                result = processor.process_image(img_path, args.confidence)
                output_file = processor.save_result_to_file(result, img_path, args.format)
                
                print(f"✓ 成功處理，識別內容:")
                if result['text_content']:
                    for line in result['text_content'].split('\n'):
                        print(f"  '{line}'")
                else:
                    print("  (無內容)")
                    
            except Exception as e:
                print(f"✗ 處理失敗: {e}")
        
        return
    
    if args.folder:
        try:
            output_files = processor.batch_process_images(args.folder, args.confidence, args.format)
            print(f"\n✓ AI-Generated Demo 批量處理完成，輸出 {len(output_files)} 個檔案")
            
        except Exception as e:
            print(f"批量處理失敗: {e}")
        return
    
    if args.file:
        try:
            result = processor.process_image(args.file, args.confidence)
            output_file = processor.save_result_to_file(result, args.file, args.format)
            
            simple_output = processor.get_config_value('OUTPUT', 'simple_output', True, bool)
            
            if not simple_output:
                print(f"\n✓ AI-Generated Demo 處理完成")
                print(f"識別內容:")
                if result['text_content']:
                    for line in result['text_content'].split('\n'):
                        print(f"  '{line}'")
                else:
                    print("  (無內容)")
                
                print(f"輸出檔案: {output_file}")
            
        except Exception as e:
            print(f"處理失敗: {e}")
        return
    
    # 互動模式
    print("=== AI-Generated PaddleOCR Testing Demo Tools ===")
    print("1. 建立並測試 AI-Generated 示例圖片")
    print("2. 處理您的圖片檔案")
    print("3. 批量處理圖片資料夾")
    print("4. 顯示當前配置")
    print("5. 退出")
    
    while True:
        choice = input("\n請選擇操作 (1-5): ").strip()
        
        if choice == '1':
            test_images = [
                ("Hello World", "ai_demo_test_en.png"),
                ("測試繁體中文識別", "ai_demo_test_cht.png"),
                ("Mixed 混合內容 123", "ai_demo_test_mixed.png")
            ]
            
            for text, filename in test_images:
                print(f"\n🔍 處理測試: {text}")
                img_path = processor.create_better_test_image(text, filename)
                
                try:
                    result = processor.process_image(img_path)
                    output_file = processor.save_result_to_file(result, img_path)
                    
                    print(f"✓ 識別內容: '{result['text_content']}'")
                    
                except Exception as e:
                    print(f"✗ 處理失敗: {e}")
        
        elif choice == '2':
            file_path = input("請輸入圖片檔案路徑: ").strip().strip('"')
            if os.path.exists(file_path):
                # 詢問是否覆蓋配置
                use_custom = input("是否自訂參數？(y/N): ").strip().lower() == 'y'
                
                if use_custom:
                    format_choice = input("選擇輸出格式 (txt/docx/pdf, 留空使用配置檔案): ").strip() or None
                    confidence_input = input("信心度閾值 (0.0-1.0, 留空使用配置檔案): ").strip()
                    confidence = float(confidence_input) if confidence_input else None
                else:
                    format_choice = None
                    confidence = None
                
                try:
                    result = processor.process_image(file_path, confidence)
                    output_file = processor.save_result_to_file(result, file_path, format_choice)
                    
                    print(f"✓ 處理完成，識別內容:")
                    if result['text_content']:
                        for line in result['text_content'].split('\n'):
                            print(f"  '{line}'")
                    else:
                        print("  (無內容)")
                    
                except Exception as e:
                    print(f"處理失敗: {e}")
            else:
                print("檔案不存在")
        
        elif choice == '3':
            folder_path = input("請輸入圖片資料夾路徑: ").strip().strip('"')
            if os.path.exists(folder_path):
                # 詢問是否覆蓋配置
                use_custom = input("是否自訂參數？(y/N): ").strip().lower() == 'y'
                
                if use_custom:
                    format_choice = input("選擇輸出格式 (txt/docx/pdf, 留空使用配置檔案): ").strip() or None
                    confidence_input = input("信心度閾值 (0.0-1.0, 留空使用配置檔案): ").strip()
                    confidence = float(confidence_input) if confidence_input else None
                else:
                    format_choice = None
                    confidence = None
                
                try:
                    output_files = processor.batch_process_images(folder_path, confidence, format_choice)
                    print(f"✓ 批量處理完成，輸出 {len(output_files)} 個檔案")
                    
                except Exception as e:
                    print(f"批量處理失敗: {e}")
            else:
                print("資料夾不存在")
        
        elif choice == '4':
            processor.print_current_config()
        
        elif choice == '5':
            print("退出 AI-Generated PaddleOCR Testing Demo Tools")
            break
        
        else:
            print("無效的選擇")


if __name__ == "__main__":
    main()